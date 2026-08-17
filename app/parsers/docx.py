#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""DOCX 解析器：python-docx 按标题样式（Heading 1/2/3）切分，段落聚合"""
from pathlib import Path
from typing import List

from parsers.base import BaseParser, Chunk


class DocxParser(BaseParser):
    format_name = "docx"

    def parse(self, path: Path) -> List[Chunk]:
        import docx  # 延迟导入

        document = docx.Document(str(path))
        chunks = []
        cur_title = ""
        cur_heading = ""
        buf = []

        def flush():
            if buf:
                text = "\n".join(buf).strip()
                if text:
                    hp = f"{cur_title} / {cur_heading}" if cur_heading else cur_title
                    chunks.append(Chunk(text=f"【{hp}】\n{text}", file=cur_title or path.stem, section=hp))
                buf.clear()

        for para in document.paragraphs:
            style = (para.style.name or "").lower()
            text = para.text.strip()
            if not text:
                continue
            if style.startswith("heading 1") or style == "title":
                flush(); cur_title = text; cur_heading = ""
            elif style.startswith("heading 2"):
                flush(); cur_heading = text
            elif style.startswith("heading 3"):
                flush(); cur_heading = f"{cur_heading} / {text}" if cur_heading else text
            else:
                buf.append(text)
        flush()

        # 无标题结构时整文件一个 chunk
        if not chunks:
            full = "\n".join(p.text.strip() for p in document.paragraphs if p.text.strip())
            if full:
                chunks.append(Chunk(text=full, file=path.stem, section=""))
        return chunks
